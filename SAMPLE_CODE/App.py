import os
import io
import json
import hashlib
import datetime
from collections import Counter

import streamlit as st
from pypdf import PdfReader
from docx import Document
from pydantic import BaseModel, Field
from google import genai
from google.genai import types
from dotenv import load_dotenv
from supabase import create_client

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

# =====================================================================
# 0. CONFIG & CLIENTS
# =====================================================================
load_dotenv()

SUPABASE_URL = os.environ.get("SUPABASE_URL")
SUPABASE_KEY = os.environ.get("SUPABASE_KEY")

@st.cache_resource
def get_supabase():
    return create_client(SUPABASE_URL, SUPABASE_KEY)

st.set_page_config(page_title="AI Contract Risk Analyzer", layout="wide")

# =====================================================================
# 1. AUTHENTICATION (Supabase-backed users table: app_users)
# =====================================================================
def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode("utf-8")).hexdigest()

def register_user(username, email, password):
    sb = get_supabase()
    existing = sb.table("app_users").select("id").eq("username", username).execute()
    if existing.data:
        return False, "Username already exists."
    sb.table("app_users").insert({
        "username": username,
        "email": email,
        "password_hash": hash_password(password),
        "role": "user",
        "created_at": datetime.datetime.utcnow().isoformat(),
    }).execute()
    return True, "Account created. Please log in."

def login_user(username, password):
    sb = get_supabase()
    result = sb.table("app_users").select("*").eq("username", username).execute()
    if not result.data:
        return None
    user = result.data[0]
    if user["password_hash"] == hash_password(password):
        return user
    return None

def render_auth_screen():
    st.title("🛡️ AI Contract Risk Analyzer")
    st.caption("Please log in or create an account to continue.")

    tab_login, tab_register = st.tabs(["🔐 Login", "🆕 Register"])

    with tab_login:
        with st.form("login_form"):
            username = st.text_input("Username")
            password = st.text_input("Password", type="password")
            submitted = st.form_submit_button("Log In")
        if submitted:
            try:
                user = login_user(username.strip(), password)
            except Exception as e:
                st.error(f"Login failed: {e}")
                user = None
            if user:
                st.session_state.auth_user = user
                st.success(f"Welcome back, {user['username']}!")
                st.rerun()
            else:
                st.error("Invalid username or password.")

    with tab_register:
        with st.form("register_form"):
            new_username = st.text_input("Choose a username")
            new_email = st.text_input("Email")
            new_password = st.text_input("Choose a password", type="password")
            reg_submitted = st.form_submit_button("Create Account")
        if reg_submitted:
            if not new_username or not new_password:
                st.error("Username and password are required.")
            else:
                try:
                    ok, msg = register_user(new_username.strip(), new_email.strip(), new_password)
                except Exception as e:
                    ok, msg = False, f"Registration failed: {e}"
                (st.success if ok else st.error)(msg)

# =====================================================================
# 2. FILE EXTRACTION PIPELINE
# =====================================================================
def extract_text_from_upload(uploaded_file):
    """Extracts raw text strings from PDF, DOCX, or TXT file bytes."""
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    try:
        if file_ext == ".txt":
            return uploaded_file.read().decode("utf-8")
        elif file_ext == ".pdf":
            reader = PdfReader(uploaded_file)
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text
        elif file_ext == ".docx":
            doc = Document(uploaded_file)
            return "\n".join([p.text for p in doc.paragraphs])
        else:
            return None
    except Exception as e:
        st.error(f"File extraction subsystem error: {e}")
        return None

# =====================================================================
# 3. ENTERPRISE PYDANTIC LEGAL SCHEMAS
# =====================================================================
class LegalRisk(BaseModel):
    risk_title: str = Field(description="The formal legal name or category of the identified liability or missing clause.")
    severity: str = Field(description="The risk classification. Must strictly map to: High, Medium, or Low.")
    confidence_score: int = Field(description="The AI model's internal confidence index regarding this risk, graded deterministically from 0 to 100.")
    explanation: str = Field(description="A comprehensive, professional legal explanation detailing exactly why this clause or omission poses a commercial threat.")

class ContractAnalysisSchema(BaseModel):
    contract_type: str = Field(description="The categorized legal type of document (e.g., NDA, Employment Contract, SaaS Agreement).")
    parties_involved: list[str] = Field(description="Full legal names of all signing parties.")
    effective_date: str = Field(description="The execution or start date. 'Not Specified' if missing.")
    expiry_date: str = Field(description="The termination, end date, or renewal checkpoint. 'Not Specified' if missing.")
    payment_terms: str = Field(description="A summary of payment amounts, schedules, and conditions found in the document. 'Not Specified' if absent.")
    renewal_clause: str = Field(description="A summary of any auto-renewal or renewal terms. 'Not Specified' if absent.")
    confidentiality_clause: str = Field(description="A summary of confidentiality/NDA obligations. 'Not Specified' if absent.")
    termination_clause: str = Field(description="A summary of how and under what conditions the agreement can be terminated. 'Not Specified' if absent.")
    responsibilities: list[str] = Field(description="Key responsibilities/obligations assigned to each party.")
    key_obligations: list[str] = Field(description="The most important ongoing obligations a signee must be aware of.")
    recommended_actions: list[str] = Field(description="Concrete, actionable recommendations for the reviewing party before signing.")
    executive_summary: str = Field(description="A concise, high-level executive summary of the agreement's purpose and scope.")
    overall_risk_score: int = Field(description="An overall document risk score from 0 (no risk) to 100 (severe risk), based on the number and severity of detected risks.")
    detected_risks: list[LegalRisk] = Field(description="A structured list of all flagged legal red flags, ambiguous terms, or liabilities.")

# =====================================================================
# 4. GEMINI LEGAL GATEWAY
# =====================================================================
def analyze_legal_document(document_text):
    client = genai.Client()
    prompt = f"""
You are an elite corporate legal counsel, chief contract auditor, and senior risk assessment manager.
Your primary task is to conduct a highly critical, deterministic audit of the provided legal document text.

Analyze the contract text thoroughly to extract administrative metadata, payment terms, renewal terms,
confidentiality terms, termination terms, responsibilities, key obligations, recommended actions before
signing, an overall risk score, and a structured list of any major liabilities, legal traps, unusual
payment terms, missing protection clauses, or ambiguous operational parameters.

Contract Text Ingestion:
---
{document_text}
---
"""
    ai_config = types.GenerateContentConfig(
        temperature=0.1,
        response_mime_type="application/json",
        response_schema=ContractAnalysisSchema,
    )
    response = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=prompt,
        config=ai_config,
    )
    return response.text

def ask_question_about_document(document_text, question):
    """Lightweight semantic Q&A over the currently loaded document (no vector DB required)."""
    client = genai.Client()
    prompt = f"""
You are a precise legal document assistant. Answer the user's question using ONLY information
found in the contract text below. If the answer is not present in the text, say so explicitly.
Quote or point to the relevant clause where possible, in your own words.

Contract Text:
---
{document_text}
---

Question: {question}

Answer concisely and specifically.
"""
    response = client.models.generate_content(model="gemini-2.5-flash", contents=prompt)
    return response.text

# =====================================================================
# 5. CLOUD AUDIT LOGGING
# =====================================================================
def sync_analysis_to_cloud(file_name, contract_data_dict, uploaded_by):
    sb = get_supabase()
    row_data = {
        "file_name": file_name,
        "contract_type": contract_data_dict.get("contract_type"),
        "effective_date": contract_data_dict.get("effective_date"),
        "expiry_date": contract_data_dict.get("expiry_date"),
        "summary": contract_data_dict.get("executive_summary"),
        "overall_risk_score": contract_data_dict.get("overall_risk_score"),
        "uploaded_by": uploaded_by,
        "raw_ai_output": contract_data_dict,
        "created_at": datetime.datetime.utcnow().isoformat(),
    }
    return sb.table("legal_contract_analyses").insert(row_data).execute()

def fetch_history(username=None, is_admin=False):
    sb = get_supabase()
    query = sb.table("legal_contract_analyses").select("*").order("created_at", desc=True)
    if not is_admin:
        query = query.eq("uploaded_by", username)
    return query.execute().data

# =====================================================================
# 6. REPORT EXPORT (PDF / DOCX)
# =====================================================================
def build_pdf_report(report: dict, file_name: str) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=50, bottomMargin=50, leftMargin=50, rightMargin=50)
    styles = getSampleStyleSheet()
    h1 = ParagraphStyle("h1c", parent=styles["Title"], spaceAfter=16)
    h2 = ParagraphStyle("h2c", parent=styles["Heading2"], spaceBefore=14, spaceAfter=6)
    body = ParagraphStyle("bodyc", parent=styles["Normal"], spaceAfter=8, leading=14)

    story = [Paragraph("AI Contract Risk Assessment Report", h1)]
    story.append(Paragraph(f"Source File: {file_name}", body))
    story.append(Paragraph(f"Generated: {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}", body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Document Metadata", h2))
    story.append(Paragraph(f"<b>Contract Type:</b> {report.get('contract_type')}", body))
    story.append(Paragraph(f"<b>Parties Involved:</b> {', '.join(report.get('parties_involved', []))}", body))
    story.append(Paragraph(f"<b>Effective Date:</b> {report.get('effective_date')}", body))
    story.append(Paragraph(f"<b>Expiry Date:</b> {report.get('expiry_date')}", body))
    story.append(Paragraph(f"<b>Overall Risk Score:</b> {report.get('overall_risk_score')}/100", body))

    story.append(Paragraph("Executive Summary", h2))
    story.append(Paragraph(report.get("executive_summary", ""), body))

    story.append(Paragraph("Key Clause Analysis", h2))
    story.append(Paragraph(f"<b>Payment Terms:</b> {report.get('payment_terms')}", body))
    story.append(Paragraph(f"<b>Renewal Clause:</b> {report.get('renewal_clause')}", body))
    story.append(Paragraph(f"<b>Confidentiality Clause:</b> {report.get('confidentiality_clause')}", body))
    story.append(Paragraph(f"<b>Termination Clause:</b> {report.get('termination_clause')}", body))

    story.append(Paragraph("Responsibilities", h2))
    for item in report.get("responsibilities", []):
        story.append(Paragraph(f"• {item}", body))

    story.append(Paragraph("Key Obligations", h2))
    for item in report.get("key_obligations", []):
        story.append(Paragraph(f"• {item}", body))

    story.append(Paragraph("Recommended Actions", h2))
    for item in report.get("recommended_actions", []):
        story.append(Paragraph(f"• {item}", body))

    story.append(Paragraph("Detected Risks", h2))
    for risk in report.get("detected_risks", []):
        story.append(Paragraph(
            f"<b>[{risk.get('severity')}] {risk.get('risk_title')}</b> — Confidence: {risk.get('confidence_score')}%",
            body,
        ))
        story.append(Paragraph(risk.get("explanation", ""), body))

    doc.build(story)
    buffer.seek(0)
    return buffer.getvalue()

def build_docx_report(report: dict, file_name: str) -> bytes:
    doc = Document()
    doc.add_heading("AI Contract Risk Assessment Report", level=0)
    doc.add_paragraph(f"Source File: {file_name}")
    doc.add_paragraph(f"Generated: {datetime.datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}")

    doc.add_heading("Document Metadata", level=1)
    doc.add_paragraph(f"Contract Type: {report.get('contract_type')}")
    doc.add_paragraph(f"Parties Involved: {', '.join(report.get('parties_involved', []))}")
    doc.add_paragraph(f"Effective Date: {report.get('effective_date')}")
    doc.add_paragraph(f"Expiry Date: {report.get('expiry_date')}")
    doc.add_paragraph(f"Overall Risk Score: {report.get('overall_risk_score')}/100")

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(report.get("executive_summary", ""))

    doc.add_heading("Key Clause Analysis", level=1)
    doc.add_paragraph(f"Payment Terms: {report.get('payment_terms')}")
    doc.add_paragraph(f"Renewal Clause: {report.get('renewal_clause')}")
    doc.add_paragraph(f"Confidentiality Clause: {report.get('confidentiality_clause')}")
    doc.add_paragraph(f"Termination Clause: {report.get('termination_clause')}")

    doc.add_heading("Responsibilities", level=1)
    for item in report.get("responsibilities", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Key Obligations", level=1)
    for item in report.get("key_obligations", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Recommended Actions", level=1)
    for item in report.get("recommended_actions", []):
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Detected Risks", level=1)
    for risk in report.get("detected_risks", []):
        doc.add_paragraph(f"[{risk.get('severity')}] {risk.get('risk_title')} — Confidence: {risk.get('confidence_score')}%", style="Heading3")
        doc.add_paragraph(risk.get("explanation", ""))

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()

# =====================================================================
# 7. UI — MAIN APP (post-login)
# =====================================================================
def render_risk_block(risk_item, idx):
    severity_str = risk_item.get("severity", "Low").lower()
    box_icon = "🔴" if "high" in severity_str else "🟡" if "medium" in severity_str else "🔵"
    with st.expander(f"{box_icon} Risk {idx + 1}: {risk_item.get('risk_title')} [Severity: {risk_item.get('severity')}]"):
        st.metric(label="AI Confidence Level Score", value=f"{risk_item.get('confidence_score')}%")
        st.markdown(f"**Legal Assessment Details:** {risk_item.get('explanation')}")

def render_analyze_tab(user):
    uploaded_contract = st.file_uploader(
        "Upload Document for Corporate Audit (.pdf, .docx, .txt)",
        type=["txt", "pdf", "docx"],
    )

    if uploaded_contract is None:
        return

    st.success(f"🔒 Secure File Received: **{uploaded_contract.name}**")

    with st.spinner("Extracting structural text matrices..."):
        contract_raw_text = extract_text_from_upload(uploaded_contract)

    if contract_raw_text is None:
        st.error("Failed to process document format.")
        return

    with st.expander("👁️ View Extracted Document Source Text"):
        st.text(contract_raw_text[:1500] + "\n\n[...Truncated for Display Preview...]")

    # --- Lightweight semantic search / Q&A over this document ---
    with st.expander("🔎 Ask a question about this document (semantic search)"):
        question = st.text_input("e.g. 'What are the payment terms?' or 'Show termination conditions.'")
        if st.button("Search Document") and question.strip():
            with st.spinner("Searching document semantically..."):
                try:
                    answer = ask_question_about_document(contract_raw_text, question)
                    st.info(answer)
                except Exception as e:
                    st.error(f"Search failed: {e}")

    if st.button("⚡ Execute AI Corporate Contract Audit"):
        with st.spinner("Analyzing clauses via Gemini Legal Intelligence Engine..."):
            try:
                raw_report_json = analyze_legal_document(contract_raw_text)
                parsed_report = json.loads(raw_report_json)
                st.session_state.last_report = parsed_report
                st.session_state.last_file_name = uploaded_contract.name
                st.success("🎯 Analysis Complete! Displaying Structured Risk Records:")
            except Exception as error:
                st.error(f"AI Pipeline Execution Interrupted: {error}")
                return

    parsed_report = st.session_state.get("last_report")
    if not parsed_report:
        return

    meta_col, summary_col = st.columns([1, 2])
    with meta_col:
        st.subheader("📋 Document Metadata")
        st.info(f"**Agreement Type:** {parsed_report.get('contract_type')}")
        st.write(f"**Effective Date:** {parsed_report.get('effective_date')}")
        st.write(f"**Expiry Date:** {parsed_report.get('expiry_date')}")
        st.metric("Overall Risk Score", f"{parsed_report.get('overall_risk_score', 0)}/100")
        st.write("**Signing Parties Found:**")
        for party in parsed_report.get("parties_involved", []):
            st.code(party)

    with summary_col:
        st.subheader("📝 Executive Summary")
        st.write(parsed_report.get("executive_summary"))

        st.subheader("📌 Key Clause Analysis")
        st.write(f"**Payment Terms:** {parsed_report.get('payment_terms')}")
        st.write(f"**Renewal Clause:** {parsed_report.get('renewal_clause')}")
        st.write(f"**Confidentiality Clause:** {parsed_report.get('confidentiality_clause')}")
        st.write(f"**Termination Clause:** {parsed_report.get('termination_clause')}")

    col_a, col_b = st.columns(2)
    with col_a:
        st.subheader("✅ Responsibilities")
        for item in parsed_report.get("responsibilities", []):
            st.write(f"- {item}")
        st.subheader("📎 Key Obligations")
        for item in parsed_report.get("key_obligations", []):
            st.write(f"- {item}")
    with col_b:
        st.subheader("💡 Recommended Actions")
        for item in parsed_report.get("recommended_actions", []):
            st.write(f"- {item}")

    st.markdown("---")
    st.subheader("🚨 Detected Risks & Legal Red Flags")
    risks_list = parsed_report.get("detected_risks", [])
    if not risks_list:
        st.success("✅ Clean Audit: No high-level liabilities or red flags discovered.")
    else:
        for idx, risk_item in enumerate(risks_list):
            render_risk_block(risk_item, idx)

    st.markdown("---")
    st.subheader("📤 Export Report")
    exp_col1, exp_col2, exp_col3 = st.columns(3)
    with exp_col1:
        pdf_bytes = build_pdf_report(parsed_report, st.session_state.last_file_name)
        st.download_button("⬇️ Download PDF Report", data=pdf_bytes,
                            file_name=f"risk_report_{st.session_state.last_file_name}.pdf",
                            mime="application/pdf")
    with exp_col2:
        docx_bytes = build_docx_report(parsed_report, st.session_state.last_file_name)
        st.download_button("⬇️ Download DOCX Report", data=docx_bytes,
                            file_name=f"risk_report_{st.session_state.last_file_name}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    with exp_col3:
        if st.button("☁️ Save to Cloud History"):
            with st.spinner("Syncing corporate records to cloud Supabase tables..."):
                try:
                    sync_analysis_to_cloud(st.session_state.last_file_name, parsed_report, user["username"])
                    st.success("☁️ Cloud Ledger Synchronization Successful. Audit complete.")
                except Exception as e:
                    st.error(f"Cloud sync failed: {e}")

def render_history_tab(user):
    st.subheader("📚 Document Analysis History")
    try:
        records = fetch_history(username=user["username"], is_admin=(user["role"] == "admin"))
    except Exception as e:
        st.error(f"Could not load history: {e}")
        return

    if not records:
        st.info("No analyses saved yet. Run an audit and click 'Save to Cloud History'.")
        return

    for rec in records:
        risk_score = rec.get("overall_risk_score", "N/A")
        with st.expander(f"📄 {rec.get('file_name')} — {rec.get('contract_type')} — Risk: {risk_score}/100 ({rec.get('created_at', '')[:10]})"):
            st.write(f"**Uploaded by:** {rec.get('uploaded_by')}")
            st.write(f"**Effective Date:** {rec.get('effective_date')}")
            st.write(f"**Expiry Date:** {rec.get('expiry_date')}")
            st.write(f"**Summary:** {rec.get('summary')}")

def render_dashboard_tab(user):
    st.subheader("📊 AI Insights Dashboard")
    try:
        records = fetch_history(username=user["username"], is_admin=(user["role"] == "admin"))
    except Exception as e:
        st.error(f"Could not load dashboard data: {e}")
        return

    if not records:
        st.info("No data yet — analyze and save some documents first.")
        return

    total_docs = len(records)
    scores = [r.get("overall_risk_score") for r in records if r.get("overall_risk_score") is not None]
    avg_risk = round(sum(scores) / len(scores), 1) if scores else 0
    high_risk_docs = sum(1 for s in scores if s >= 70)

    risk_titles = []
    for r in records:
        raw = r.get("raw_ai_output") or {}
        for risk in raw.get("detected_risks", []):
            risk_titles.append(risk.get("risk_title"))
    common_risks = Counter(risk_titles).most_common(5)

    c1, c2, c3 = st.columns(3)
    c1.metric("Total Documents Analyzed", total_docs)
    c2.metric("Average Risk Score", f"{avg_risk}/100")
    c3.metric("High-Risk Documents (≥70)", high_risk_docs)

    st.markdown("#### 🔁 Most Frequently Detected Risks")
    if common_risks:
        for title, count in common_risks:
            st.write(f"- **{title}** — flagged {count}x")
    else:
        st.write("No risks recorded yet.")

def render_admin_tab():
    st.subheader("🛠️ Admin Panel")
    sb = get_supabase()

    st.markdown("#### 👥 Registered Users")
    try:
        users = sb.table("app_users").select("username, email, role, created_at").execute().data
        st.dataframe(users, use_container_width=True)
    except Exception as e:
        st.error(f"Could not load users: {e}")

    st.markdown("#### 📈 System-Wide Processing Stats")
    try:
        all_records = sb.table("legal_contract_analyses").select("*").execute().data
        st.metric("Total Documents Processed (All Users)", len(all_records))
        scores = [r.get("overall_risk_score") for r in all_records if r.get("overall_risk_score") is not None]
        if scores:
            st.metric("System-Wide Average Risk Score", f"{round(sum(scores)/len(scores), 1)}/100")
    except Exception as e:
        st.error(f"Could not load system stats: {e}")

    st.markdown("#### 🗂️ Recent System Logs")
    try:
        recent = sb.table("legal_contract_analyses").select("file_name, uploaded_by, created_at").order("created_at", desc=True).limit(20).execute().data
        st.dataframe(recent, use_container_width=True)
    except Exception as e:
        st.error(f"Could not load logs: {e}")

def render_profile_sidebar(user):
    with st.sidebar:
        st.markdown(f"### 👤 {user['username']}")
        st.caption(f"Role: {user.get('role', 'user')}")
        st.caption(f"Email: {user.get('email', 'N/A')}")
        if st.button("🚪 Log Out"):
            st.session_state.auth_user = None
            st.session_state.last_report = None
            st.rerun()

# =====================================================================
# 8. APP ENTRYPOINT
# =====================================================================
def main():
    if "auth_user" not in st.session_state:
        st.session_state.auth_user = None

    if not st.session_state.auth_user:
        render_auth_screen()
        return

    user = st.session_state.auth_user
    render_profile_sidebar(user)

    st.title("🛡️ Production AI Legal Contract & Risk Analyzer")
    st.markdown("---")

    tabs = ["⚡ Analyze", "📚 History", "📊 Dashboard"]
    if user.get("role") == "admin":
        tabs.append("🛠️ Admin")

    tab_objs = st.tabs(tabs)

    with tab_objs[0]:
        render_analyze_tab(user)
    with tab_objs[1]:
        render_history_tab(user)
    with tab_objs[2]:
        render_dashboard_tab(user)
    if user.get("role") == "admin":
        with tab_objs[3]:
            render_admin_tab()

if __name__ == "__main__":
    main()